"""
Graph node functions — one function per node in the research graph.

Node execution order (happy path)
──────────────────────────────────
plan_research → human_review → execute_tasks → evaluate_results
    → synthesize_report → citation_agent → save_output

Conditional branches
────────────────────
• human_review      → plan_research        (user rejected / modified)
• execute_tasks     → handle_escalation    (tasks exhausted retries)
• evaluate_results  → plan_research        (more research needed, iter < 5)

Each node receives the full ResearchState and returns a dict containing ONLY
the fields it modifies. LangGraph merges the return dict into the state.
Fields with operator.add reducers (sub_agent_results, all_sources, error_log)
are APPENDED to automatically — nodes just return new items in a list.
"""

import json
import logging
import os
import time
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path

from docx import Document
from langchain_openai import ChatOpenAI
from langgraph.types import interrupt

from config.config import (
    OPENAI_API_KEY,
    OPENAI_MODEL,
    MAX_ITERATIONS,
    MAX_RETRY_COUNT,
    MAX_SUB_AGENTS,
    OUTPUT_DIR,
    TAVILY_API_KEY,
)
from data.research_context.session_store import SessionStore
from src.app.graph.prompts import (
    CITATION_AGENT_PROMPT,
    EVALUATE_RESULTS_PROMPT,
    HANDLE_ESCALATION_PROMPT,
    PLAN_RESEARCH_PROMPT,
    SUB_AGENT_PROMPT,
    SYNTHESIZE_REPORT_PROMPT,
)
from src.app.graph.state import ResearchState
from src.utils.tools.search import run_tavily_search

logger = logging.getLogger(__name__)

# ── Module-level LLM instance (shared across all nodes) ───────────────────────
# Initialised once to avoid re-creating the client on every node invocation.
llm = ChatOpenAI(api_key=OPENAI_API_KEY, model=OPENAI_MODEL)

# ── Module-level session store ─────────────────────────────────────────────────
session_store = SessionStore()


# ─────────────────────────────────────────────────────────────────────────────
# Helper — safe JSON parse
# ─────────────────────────────────────────────────────────────────────────────
def _parse_json(raw: str, context: str) -> dict | list:
    """Strip optional markdown fences then parse JSON.
    Raises ValueError with context info on failure."""
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        # Remove ```json ... ``` or ``` ... ```
        cleaned = cleaned.split("\n", 1)[-1]
        cleaned = cleaned.rsplit("```", 1)[0].strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error("JSON parse failed in %s: %s\nRaw: %s", context, exc, raw[:300])
        raise ValueError(f"Invalid JSON from LLM in {context}: {exc}") from exc


# ─────────────────────────────────────────────────────────────────────────────
# Node 1 — plan_research
# ─────────────────────────────────────────────────────────────────────────────
def plan_research_node(state: ResearchState) -> dict:
    """Lead Researcher: Think (Plan Approach).

    Decomposes the research goal into a JSON task list, increments the
    iteration counter, saves the plan to SQLite, and resets the human
    approval flag ready for the next human_review interrupt.
    """
    logger.info("Node: plan_research | iteration=%d", state.get("iteration_count", 0))

    research_goal: str = state["research_goal"]
    iteration: int = state.get("iteration_count", 0)
    human_feedback: str = state.get("human_feedback", "")

    # Build prior-context summary from previously completed sub-agent results
    prior_results = state.get("sub_agent_results", [])
    if prior_results:
        prior_context = json.dumps(
            [
                {"task_id": r["task_id"], "title_hint": r.get("findings", "")[:120]}
                for r in prior_results
                if r.get("status") == "completed"
            ],
            indent=2,
        )
    else:
        prior_context = ""

    # Invoke Lead Researcher LLM
    chain = PLAN_RESEARCH_PROMPT | llm
    response = chain.invoke(
        {
            "research_goal": research_goal,
            "iteration_count": iteration,
            "human_feedback": human_feedback,
            "prior_context": prior_context,
        }
    )

    tasks: list[dict] = _parse_json(response.content, "plan_research")

    # Assign UUIDs to any task missing a task_id
    for task in tasks:
        if not task.get("task_id"):
            task["task_id"] = str(uuid.uuid4())
        task.setdefault("status", "pending")
        task.setdefault("retry_count", 0)
        task.setdefault("assigned_to", "")

    new_iteration = iteration + 1

    # Persist plan to app-level SQLite store
    session_store.save_plan(
        session_id=state.get("research_goal", "default"),
        iteration=new_iteration,
        tasks=tasks,
    )

    logger.info("plan_research: %d tasks generated for iteration %d", len(tasks), new_iteration)

    return {
        "tasks": tasks,
        "iteration_count": new_iteration,
        "human_approved": False,
        "human_feedback": "",
    }


# ─────────────────────────────────────────────────────────────────────────────
# Node 2 — human_review
# ─────────────────────────────────────────────────────────────────────────────
def human_review_node(state: ResearchState) -> dict:
    """Human-in-the-Loop interrupt.

    Only interrupts on iteration 1 (the initial plan). Subsequent
    iterations are auto-approved so the research loop runs unattended.

    Expected resume payload (iteration 1 only):
        {"action": "approve"}
        {"action": "modify",  "feedback": "<instructions>"}
        {"action": "reject",  "feedback": "<instructions>"}
    """
    iteration: int = state.get("iteration_count", 1)

    # Iterations 2+ — auto-approve and continue without interrupting the user
    if iteration > 1:
        logger.info(
            "human_review: iteration %d — auto-approving (review only on iteration 1)",
            iteration,
        )
        return {"human_approved": True, "human_feedback": ""}

    logger.info("Node: human_review — pausing for human approval (iteration 1)")

    # Pause graph; return task list to the API caller for display
    response: dict = interrupt(
        {
            "message": "Please review the research task list and respond with "
                       "approve / modify / reject.",
            "tasks": state["tasks"],
            "iteration": iteration,
        }
    )

    action: str = (response or {}).get("action", "approve").lower()
    feedback: str = (response or {}).get("feedback", "")

    if action == "approve":
        logger.info("human_review: approved")
        return {"human_approved": True, "human_feedback": ""}

    # "modify" or "reject" both route back to plan_research
    logger.info("human_review: action=%s — routing back to plan_research", action)
    return {
        "human_approved": False,
        "human_feedback": feedback,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Sub-agent worker (called inside ThreadPoolExecutor)
# ─────────────────────────────────────────────────────────────────────────────
def _sub_agent_worker(
    task: dict,
    research_goal: str,
    user_documents: list[str],
) -> dict:
    """Execute one research task with retry logic.

    Steps per attempt:
      1. Run all search_queries via Tavily.
      2. Load any relevant user documents.
      3. Call LLM to synthesise findings from raw results.
      4. Return structured result dict.

    Returns an escalated result dict after MAX_RETRY_COUNT failures.
    """
    task_id: str = task["task_id"]
    last_error: str = ""

    for attempt in range(MAX_RETRY_COUNT + 1):
        try:
            # Step 1 — web search
            raw_search_results: list[dict] = []
            for query in task.get("search_queries", []):
                hits = run_tavily_search(query)
                raw_search_results.extend(hits)

            # Step 2 — relevant document content
            doc_content = ""
            relevant_docs = task.get("relevant_documents", [])
            if relevant_docs:
                from src.utils.tools.document_loader import load_document
                for doc_path in relevant_docs:
                    if doc_path in user_documents:
                        doc_content += load_document(doc_path) + "\n\n"

            # Step 3 — LLM synthesis
            chain = SUB_AGENT_PROMPT | llm
            response = chain.invoke(
                {
                    "research_goal": research_goal,
                    "task_title": task.get("title", ""),
                    "task_description": task.get("description", ""),
                    "search_results": json.dumps(raw_search_results, indent=2),
                    "document_content": doc_content.strip() or "None provided.",
                }
            )

            parsed: dict = _parse_json(response.content, f"sub_agent:{task_id}")

            logger.info(
                "sub_agent: task_id=%s completed on attempt %d", task_id, attempt + 1
            )
            return {
                "task_id": task_id,
                "findings": parsed.get("findings", ""),
                "sources": parsed.get("sources", []),
                "status": "completed",
                "retry_count": attempt,
                "error_message": None,
            }

        except Exception as exc:
            last_error = str(exc)
            logger.warning(
                "sub_agent: task_id=%s attempt %d failed: %s", task_id, attempt + 1, last_error
            )
            if attempt < MAX_RETRY_COUNT:
                time.sleep(2 ** attempt)  # exponential back-off: 1s, 2s, 4s

    # All retries exhausted — escalate
    logger.error("sub_agent: task_id=%s escalated after %d attempts", task_id, MAX_RETRY_COUNT + 1)
    return {
        "task_id": task_id,
        "findings": "",
        "sources": [],
        "status": "escalated",
        "retry_count": MAX_RETRY_COUNT,
        "error_message": last_error,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Node 3 — execute_tasks
# ─────────────────────────────────────────────────────────────────────────────
def execute_tasks_node(state: ResearchState) -> dict:
    """Run the task queue across a fixed pool of MAX_SUB_AGENTS (3) workers.

    Tasks are dispatched to whichever worker is free — dynamic assignment,
    not pre-allocated. Results are collected once all workers are idle.
    """
    logger.info(
        "Node: execute_tasks | %d tasks | %d workers",
        len(state["tasks"]),
        MAX_SUB_AGENTS,
    )

    tasks: list[dict] = state["tasks"]
    research_goal: str = state["research_goal"]
    user_documents: list[str] = state.get("user_documents", [])

    new_results: list[dict] = []
    new_sources: list[str] = []
    escalated: list[dict] = []

    with ThreadPoolExecutor(max_workers=MAX_SUB_AGENTS) as pool:
        future_to_task = {
            pool.submit(_sub_agent_worker, task, research_goal, user_documents): task
            for task in tasks
        }

        for future in as_completed(future_to_task):
            result = future.result()
            new_results.append(result)
            new_sources.extend(result.get("sources", []))

            if result["status"] == "escalated":
                escalated.append(
                    {k: v for k, v in future_to_task[future].items()}
                    | {"error_message": result["error_message"]}
                )

    # Persist results to app-level store for cross-iteration retrieval
    session_store.save_results(
        session_id=state["research_goal"],
        iteration=state.get("iteration_count", 1),
        results=new_results,
    )

    logger.info(
        "execute_tasks: %d completed, %d escalated",
        sum(1 for r in new_results if r["status"] == "completed"),
        len(escalated),
    )

    return {
        "sub_agent_results": new_results,   # appended via operator.add
        "all_sources": list(set(new_sources)),  # appended via operator.add
        "escalated_tasks": escalated,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Node 4 — handle_escalation
# ─────────────────────────────────────────────────────────────────────────────
def handle_escalation_node(state: ResearchState) -> dict:
    """Lead Researcher decides to reassign or terminate each escalated task."""
    logger.info(
        "Node: handle_escalation | %d escalated tasks", len(state["escalated_tasks"])
    )

    chain = HANDLE_ESCALATION_PROMPT | llm
    response = chain.invoke(
        {
            "research_goal": state["research_goal"],
            "escalated_tasks": json.dumps(state["escalated_tasks"], indent=2),
        }
    )

    decisions: list[dict] = _parse_json(response.content, "handle_escalation")

    updated_tasks: list[dict] = list(state.get("tasks", []))
    new_error_entries: list[str] = []

    for decision in decisions:
        task_id = decision.get("task_id")
        action = decision.get("action", "terminate")

        if action == "reassign":
            # Build a fresh task dict based on the original, with new instructions
            reassigned = {
                "task_id": task_id,
                "title": next(
                    (t["title"] for t in state["escalated_tasks"] if t["task_id"] == task_id),
                    "Reassigned task",
                ),
                "description": decision.get("new_description", ""),
                "search_queries": decision.get("new_search_queries", []),
                "relevant_documents": [],
                "priority": "high",
                "status": "pending",
                "retry_count": 0,
                "assigned_to": "",
            }
            updated_tasks.append(reassigned)
            logger.info("handle_escalation: task_id=%s reassigned", task_id)

        else:  # terminate
            reason = decision.get("reason", "No reason provided.")
            entry = f"[TERMINATED] task_id={task_id} — {reason}"
            new_error_entries.append(entry)
            logger.warning("handle_escalation: task_id=%s terminated: %s", task_id, reason)

    return {
        "tasks": updated_tasks,
        "escalated_tasks": [],           # cleared after processing
        "error_log": new_error_entries,  # appended via operator.add
    }


# ─────────────────────────────────────────────────────────────────────────────
# Node 5 — evaluate_results
# ─────────────────────────────────────────────────────────────────────────────
def evaluate_results_node(state: ResearchState) -> dict:
    """Lead Researcher: Think (Evaluate).

    Retrieves all accumulated results from the session store, assesses
    quality and completeness, and sets the research_complete flag.
    """
    logger.info(
        "Node: evaluate_results | iteration=%d", state.get("iteration_count", 1)
    )

    # Pull all results across every iteration from the store
    all_results = session_store.get_all_results(session_id=state["research_goal"])
    if not all_results:
        all_results = state.get("sub_agent_results", [])

    iteration = state.get("iteration_count", 1)

    chain = EVALUATE_RESULTS_PROMPT | llm
    response = chain.invoke(
        {
            "research_goal": state["research_goal"],
            "iteration_count": iteration,
            "max_iterations": MAX_ITERATIONS,
            "results_json": json.dumps(all_results, indent=2),
        }
    )

    verdict: dict = _parse_json(response.content, "evaluate_results")

    research_complete: bool = bool(verdict.get("research_complete", False))
    evaluation_notes: str = verdict.get("evaluation_notes", "")

    # Force completion if iteration cap reached
    if iteration >= MAX_ITERATIONS and not research_complete:
        research_complete = True
        cap_msg = (
            f"[AUTO-COMPLETE] Iteration cap ({MAX_ITERATIONS}) reached. "
            "Forcing synthesis."
        )
        evaluation_notes += f"\n{cap_msg}"
        logger.warning(cap_msg)

    logger.info(
        "evaluate_results: research_complete=%s", research_complete
    )

    return {
        "evaluation_notes": evaluation_notes,
        "research_complete": research_complete,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Node 6 — synthesize_report
# ─────────────────────────────────────────────────────────────────────────────
def synthesize_report_node(state: ResearchState) -> dict:
    """Lead Researcher: Think (Synthesize Results).

    Combines all sub-agent findings across every iteration into a single
    cohesive report with [CITE] placeholders for the Citation Agent.
    """
    logger.info("Node: synthesize_report")

    all_results = session_store.get_all_results(session_id=state["research_goal"])
    if not all_results:
        all_results = state.get("sub_agent_results", [])

    chain = SYNTHESIZE_REPORT_PROMPT | llm
    response = chain.invoke(
        {
            "research_goal": state["research_goal"],
            "evaluation_notes": state.get("evaluation_notes", ""),
            "results_json": json.dumps(all_results, indent=2),
        }
    )

    synthesized_report: str = response.content.strip()

    # Persist to session store for crash recovery
    session_store.save_synthesized_report(
        session_id=state["research_goal"],
        report=synthesized_report,
    )

    logger.info("synthesize_report: report length=%d chars", len(synthesized_report))

    return {"synthesized_report": synthesized_report}


# ─────────────────────────────────────────────────────────────────────────────
# Node 7 — citation_agent
# ─────────────────────────────────────────────────────────────────────────────
def citation_agent_node(state: ResearchState) -> dict:
    """Citation Agent: replace [CITE] markers with numbered references."""
    logger.info("Node: citation_agent")

    all_sources: list[str] = state.get("all_sources", [])
    # Deduplicate while preserving order
    seen: set = set()
    unique_sources: list[str] = []
    for s in all_sources:
        if s not in seen:
            seen.add(s)
            unique_sources.append(s)

    sources_text = "\n".join(
        f"{i + 1}. {src}" for i, src in enumerate(unique_sources)
    )

    chain = CITATION_AGENT_PROMPT | llm
    response = chain.invoke(
        {
            "synthesized_report": state["synthesized_report"],
            "sources": sources_text or "No sources collected.",
        }
    )

    final_report: str = response.content.strip()
    logger.info("citation_agent: final report length=%d chars", len(final_report))

    return {"final_report": final_report}


# ─────────────────────────────────────────────────────────────────────────────
# Node 8 — save_output
# ─────────────────────────────────────────────────────────────────────────────
def save_output_node(state: ResearchState) -> dict:
    """Render the final report as a .docx file and persist to SQLite."""
    logger.info("Node: save_output")

    final_report: str = state["final_report"]
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = Path(OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = str(output_dir / f"research_report_{timestamp}.docx")

    # Build .docx
    doc = Document()
    doc.add_heading("Research Report", level=0)

    for line in final_report.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        # Simple heading detection: lines starting with # (markdown-style)
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            doc.add_paragraph(stripped)

    doc.save(output_path)

    # Persist final output to session store
    session_store.save_final_output(
        session_id=state["research_goal"],
        final_report=final_report,
        output_path=output_path,
    )

    logger.info("save_output: report saved to %s", output_path)

    return {"output_path": output_path}
